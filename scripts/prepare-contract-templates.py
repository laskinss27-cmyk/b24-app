from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ASSETS = Path(__file__).resolve().parents[1] / "packages" / "backend" / "assets"


def set_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    set_paragraph(paragraph, paragraph.text.replace(old, new))


def mark_requisites_table(table) -> None:
    set_paragraph(table.cell(0, 0).paragraphs[0], "{{CONTRACTOR_REQUISITES}}")
    set_paragraph(table.cell(0, 1).paragraphs[0], "{{CUSTOMER_REQUISITES}}")


def mark_signature_table(table) -> None:
    set_paragraph(table.cell(0, 0).paragraphs[0], "{{CONTRACTOR_SIGNATURE}}")
    set_paragraph(table.cell(0, 1).paragraphs[0], "{{CUSTOMER_SIGNATURE}}")


def remove_highlights(document) -> None:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = None


def document_paragraphs(document):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def replace_everywhere(document, old: str, new: str) -> None:
    for paragraph in document_paragraphs(document):
        replace_in_paragraph(paragraph, old, new)


def unique_row_cells(row):
    result = []
    seen = set()
    for cell in row.cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        result.append(cell)
    return result


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def prepare_supply() -> None:
    path = ASSETS / "contract-supply.docx"
    document = Document(path)

    corrections = {
        "5.3.4. в момент сдачи": "5.3.3. в момент сдачи",
        "участия в составления соответствующего Акта": "участия в составлении соответствующего Акта",
        "указывать кода (артикула) Товара": "указывать код (артикул) Товара",
        "ремонт осуществляются на платной основе": "ремонт осуществляется на платной основе",
        "Стороны преследует исключительно деловые цели": "Стороны преследуют исключительно деловые цели",
        "Поставщик не несет ответственность за нарушение сроков поставки": "Поставщик не несет ответственности за нарушение сроков поставки",
        "указанный в разделе 13 настоящего Договора": "указанный в разделе 14 настоящего Договора",
        "№ 63 ФЗ «Об электронной подписи»": "№ 63-ФЗ «Об электронной подписи»",
        "в соответствии с нормами Федеральным законом": "в соответствии с нормами Федерального закона",
    }
    for old, new in corrections.items():
        replace_everywhere(document, old, new)

    set_paragraph(document.paragraphs[0], "ДОГОВОР ПОСТАВКИ № {{CONTRACT_NUMBER}}")
    set_paragraph(document.paragraphs[2], "{{CONTRACTOR_PREAMBLE}}")
    set_paragraph(document.paragraphs[3], "{{CUSTOMER_PREAMBLE}}")
    header = document.tables[0]
    set_paragraph(header.cell(0, 0).paragraphs[0], "{{CITY}}")
    set_paragraph(header.cell(0, 1).paragraphs[0], "{{CONTRACT_DATE}}")

    for paragraph in document_paragraphs(document):
        replace_in_paragraph(paragraph, "buh@homelogicsoft.com", "{{CONTRACTOR_EMAIL}}")

    mark_requisites_table(document.tables[1])

    set_paragraph(
        document.paragraphs[153],
        "к Договору поставки № {{CONTRACT_NUMBER}} от {{CONTRACT_DATE}}",
    )
    set_paragraph(document.paragraphs[155], "СПЕЦИФИКАЦИЯ № 1 (к отдельной поставке)")
    set_paragraph(document.paragraphs[157], "{{CITY}}\t{{CONTRACT_DATE}}")
    set_paragraph(document.paragraphs[158], "{{CONTRACTOR_SPEC_PREAMBLE}}")
    set_paragraph(document.paragraphs[159], "{{CUSTOMER_SPEC_PREAMBLE}}")

    products = document.tables[2]
    while len(products.rows) > 3:
        products._tbl.remove(products.rows[2]._tr)
    product_cells = unique_row_cells(products.rows[1])
    product_values = (
        "{{PRODUCT_INDEX}}",
        "{{PRODUCT_NAME}}",
        "{{PRODUCT_UNIT}}",
        "{{PRODUCT_PRICE}}",
        "{{PRODUCT_QTY}}",
        "{{PRODUCT_TOTAL}}",
    )
    if len(product_cells) != len(product_values):
        raise RuntimeError(f"unexpected supply product columns: {len(product_cells)}")
    for cell, value in zip(product_cells, product_values):
        set_paragraph(cell.paragraphs[0], value)
    set_paragraph(unique_row_cells(products.rows[-1])[-1].paragraphs[0], "{{TOTAL}}")

    set_paragraph(
        document.paragraphs[163],
        "Общая стоимость оборудования по настоящей Спецификации составляет "
        "{{TOTAL}} ({{TOTAL_WORDS}}), включая НДС {{VAT_RATE}}%.",
    )
    set_paragraph(
        document.paragraphs[164],
        "Предоплата в размере {{SUPPLY_PREPAYMENT_PERCENT}}% составляет "
        "{{ADVANCE}} ({{ADVANCE_WORDS}}), включая НДС {{VAT_RATE}}%.",
    )
    set_paragraph(
        document.paragraphs[165],
        "Срок передачи Товара: {{SUPPLY_DELIVERY_DURATION}} с момента оплаты предоплаты "
        "и заключения Договора.",
    )
    mark_requisites_table(document.tables[3])
    remove_paragraph(document.paragraphs[166])
    remove_highlights(document)
    document.save(path)


def prepare_design() -> None:
    path = ASSETS / "contract-design.docx"
    document = Document(path)
    set_paragraph(document.paragraphs[0], "Договор № {{CONTRACT_NUMBER}}")
    set_paragraph(document.paragraphs[4], "{{CITY}}\t{{CONTRACT_DATE}}")
    set_paragraph(document.paragraphs[6], "{{CONTRACTOR_PREAMBLE}}")
    set_paragraph(document.paragraphs[7], "{{CUSTOMER_PREAMBLE}}")
    replace_in_paragraph(
        document.paragraphs[11],
        "[наименование, адрес, описание объекта]",
        "{{OBJECT_NAME}}, адрес: {{OBJECT_ADDRESS}}",
    )
    for index in (157, 199, 218, 233):
        text = document.paragraphs[index].text
        text = text.replace("536", "{{CONTRACT_NUMBER}}")
        if "от" in text:
            text = text.split(" от", 1)[0] + " от {{CONTRACT_DATE}}"
        set_paragraph(document.paragraphs[index], text)
    set_paragraph(
        document.paragraphs[225],
        "Общая стоимость работ по настоящему Договору составляет "
        "{{TOTAL}} ({{TOTAL_WORDS}}), включая НДС {{VAT_RATE}}%.",
    )
    set_paragraph(
        document.paragraphs[227],
        "1. Заказчик производит оплату аванса в размере "
        "{{ADVANCE}} ({{ADVANCE_WORDS}}), включая НДС {{VAT_RATE}}% при подписании настоящего Договора.",
    )
    set_paragraph(
        document.paragraphs[228],
        "2. Оставшаяся сумма {{BALANCE}} ({{BALANCE_WORDS}}), включая НДС {{VAT_RATE}}% "
        "выплачивается после подписания акта приема-передачи документации "
        "и согласования Исполнителем проектов с Заказчиком.",
    )
    mark_requisites_table(document.tables[0])
    for index in (1, 3, 5, 6):
        mark_signature_table(document.tables[index])
    remove_highlights(document)
    document.save(path)


def prepare_smart_home() -> None:
    path = ASSETS / "contract-smart-home.docx"
    document = Document(path)
    set_paragraph(document.paragraphs[0], "Договор подряда № {{CONTRACT_NUMBER}}")
    set_paragraph(document.paragraphs[2], "{{CONTRACTOR_PREAMBLE}}")
    set_paragraph(document.paragraphs[3], "{{CUSTOMER_PREAMBLE}}")
    set_paragraph(
        document.paragraphs[11],
        "2.1.1. Выполнить свои обязательства в полном объеме в соответствии с "
        "Приложениями к Договору в сроки, указанные в п. 3.1. Договора.",
    )
    replace_in_paragraph(
        document.paragraphs[79],
        "buh@homelogicsoft.com",
        "{{CONTRACTOR_EMAIL}}",
    )
    set_paragraph(
        document.paragraphs[97],
        "к Договору подряда № {{CONTRACT_NUMBER}} от {{CONTRACT_DATE}}",
    )
    set_paragraph(document.paragraphs[99], "{{CITY}}\t{{CONTRACT_DATE}}")
    set_paragraph(document.paragraphs[101], "{{CONTRACTOR_AGREEMENT_PREAMBLE}}")
    set_paragraph(document.paragraphs[102], "{{CUSTOMER_AGREEMENT_PREAMBLE}}")
    set_paragraph(
        document.paragraphs[105],
        "1.1. Стороны согласовали выполнение работ по Договору подряда "
        "№ {{CONTRACT_NUMBER}} от {{CONTRACT_DATE}} (далее — «Договор»).",
    )
    set_paragraph(
        document.paragraphs[128],
        "к Дополнительному соглашению № 1 от {{CONTRACT_DATE}}",
    )
    set_paragraph(
        document.paragraphs[129],
        "к Договору подряда № {{CONTRACT_NUMBER}} от {{CONTRACT_DATE}}",
    )
    set_paragraph(document.paragraphs[136], "1.2 Адрес объекта: {{OBJECT_ADDRESS}}")
    set_paragraph(
        document.paragraphs[140],
        "2.1 Срок выполнения работ: {{WORK_DURATION}} с момента выполнения Заказчиком "
        "условий, предусмотренных п. 2.3.5 и п. 4.3 Договора.",
    )

    header = document.tables[0]
    set_paragraph(header.cell(0, 0).paragraphs[0], "{{CITY}}")
    set_paragraph(header.cell(0, 1).paragraphs[0], "{{CONTRACT_DATE}}")
    mark_requisites_table(document.tables[1])
    mark_signature_table(document.tables[2])
    mark_signature_table(document.tables[4])

    works = document.tables[3]
    while len(works.rows) > 3:
        works._tbl.remove(works.rows[2]._tr)
    row = works.rows[1]
    values = (
        "{{PRODUCT_INDEX}}",
        "{{PRODUCT_NAME}}",
        "шт.",
        "{{PRODUCT_QTY}}",
        "{{PRODUCT_TOTAL}}",
    )
    for cell, value in zip(row.cells, values):
        set_paragraph(cell.paragraphs[0], value)
    total_row = works.rows[-1]
    set_paragraph(total_row.cells[-1].paragraphs[0], "{{TOTAL}}")

    remove_highlights(document)
    document.save(path)


if __name__ == "__main__":
    prepare_supply()
    prepare_design()
    prepare_smart_home()
